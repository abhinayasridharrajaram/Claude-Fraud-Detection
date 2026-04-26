"""
Convert outputs/report.md -> outputs/USAA_Fraud_Portfolio.docx
Calibri 11, 1-inch margins, bold section headers.
Images inserted at specific section boundaries.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

BASE = os.path.dirname(os.path.abspath(__file__))
MD   = os.path.join(BASE, "outputs", "report.md")
OUT  = os.path.join(BASE, "outputs", "USAA_Fraud_Portfolio.docx")
IMG_PR   = os.path.join(BASE, "outputs", "ml", "anomaly_pr_curves.png")
IMG_UMAP = os.path.join(BASE, "outputs", "ml", "fraud_overlay_umap.png")

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, italic=False, color=None, mono=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Courier New" if mono else "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)

def body_para(doc, text="", bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p

def heading(doc, text, level):
    """Bold Calibri heading without using Word's built-in Heading styles."""
    sizes = {1: 16, 2: 13, 3: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, 11), bold=True)
    # add 6pt space before
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_image(doc, path, caption=None, width=Inches(5.5)):
    doc.add_picture(path, width=width)
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = body_para(doc, caption, italic=True)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)

def add_hrule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)

# ── table parser ─────────────────────────────────────────────────────────────

def is_separator(line):
    return bool(re.match(r"^\|[-| :]+\|$", line.strip()))

def parse_md_table(lines):
    """Return list-of-lists from markdown table lines (incl. separator)."""
    rows = []
    for ln in lines:
        if is_separator(ln):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    return rows

def add_md_table(doc, rows):
    if not rows:
        return
    ncols = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            # strip bold markers for display; re-apply as run bold
            text = cell_text.replace("**", "")
            is_bold = "**" in cell_text or r_idx == 0
            run = para.add_run(text)
            set_font(run, size=10, bold=is_bold)
    doc.add_paragraph()  # spacing after table

# ── inline markdown → runs ────────────────────────────────────────────────────

def add_inline(para, text):
    """Parse **bold**, `code`, and plain text into runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            set_font(run, mono=True)
        else:
            run = para.add_run(part)
            set_font(run)

# ── section/image markers ─────────────────────────────────────────────────────
# We'll track these flags while streaming lines.
INSERT_PR_AFTER   = "### 4.3 Performance Metrics"
INSERT_UMAP_AFTER = "## 5. Clustering Insights"   # end of section 5

# ── main converter ────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # ── page margins (1 inch) ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── default paragraph font ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with open(MD, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    # flags for image insertion
    pr_inserted   = False
    umap_inserted = False
    # track last H2 heading text so we know when section 5 ends
    in_section_5  = False
    table_lines   = []
    in_table      = False

    while i < len(lines):
        raw  = lines[i].rstrip("\n")
        line = raw.strip()

        # ── horizontal rule ──
        if line in ("---", "___", "***"):
            if in_table:
                add_md_table(doc, parse_md_table(table_lines))
                table_lines = []
                in_table = False
            add_hrule(doc)
            i += 1
            continue

        # ── headings ──
        h1 = re.match(r"^# (.+)$", line)
        h2 = re.match(r"^## (.+)$", line)
        h3 = re.match(r"^### (.+)$", line)

        if h1 or h2 or h3:
            if in_table:
                add_md_table(doc, parse_md_table(table_lines))
                table_lines = []
                in_table = False

            # ── image: after section 4.3 table, before 4.4 ──
            if h3 and h3.group(1).startswith("4.4") and not pr_inserted:
                add_image(doc, IMG_PR,
                          caption="Figure 1 — Precision-Recall curves for Isolation Forest, One-Class SVM, and LOF.")
                pr_inserted = True

            # ── detect end of section 5 (when we hit ## 6) ──
            if h2 and h2.group(1).startswith("6.") and in_section_5 and not umap_inserted:
                add_image(doc, IMG_UMAP,
                          caption="Figure 2 — UMAP projection colored by Class (fraud = red, legitimate = gray).")
                umap_inserted = True
                in_section_5 = False

            if h2 and h2.group(1).startswith("5."):
                in_section_5 = True

            lvl  = 1 if h1 else (2 if h2 else 3)
            text = (h1 or h2 or h3).group(1)
            heading(doc, text, lvl)
            i += 1
            continue

        # ── table ──
        if line.startswith("|"):
            if not in_table:
                in_table    = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                add_md_table(doc, parse_md_table(table_lines))
                table_lines = []
                in_table = False

        # ── blank line ──
        if not line:
            doc.add_paragraph()
            i += 1
            continue

        # ── bullet list ──
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            for run in p.runs:
                run.font.size = Pt(11)
            i += 1
            continue

        # ── numbered list ──
        nm = re.match(r"^\d+\. (.+)$", line)
        if nm:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, nm.group(1))
            for run in p.runs:
                run.font.size = Pt(11)
            i += 1
            continue

        # ── block quote / metadata lines (> ...) ──
        if line.startswith("> "):
            p = body_para(doc, italic=True)
            add_inline(p, line[2:])
            i += 1
            continue

        # ── skip TOC anchor lines ──
        if re.match(r"^\d+\. \[", line):
            i += 1
            continue

        # ── regular paragraph ──
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    # flush any dangling table
    if in_table:
        add_md_table(doc, parse_md_table(table_lines))

    # safety: insert images if markers were never triggered
    if not pr_inserted:
        add_image(doc, IMG_PR,
                  caption="Figure 1 — Precision-Recall curves for Isolation Forest, One-Class SVM, and LOF.")
    if not umap_inserted:
        add_image(doc, IMG_UMAP,
                  caption="Figure 2 — UMAP projection colored by Class (fraud = red, legitimate = gray).")

    doc.save(OUT)
    print(f"Saved: {OUT}")

if __name__ == "__main__":
    build_doc()
